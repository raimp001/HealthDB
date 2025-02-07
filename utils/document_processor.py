"""
Document processing utilities for handling various file formats
"""
import pandas as pd
from pathlib import Path
import docx
from pdf2image import convert_from_bytes
import io
import nltk
from nltk.tokenize import sent_tokenize

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

def process_text_chunks(text: str, chunk_size: int = 1000) -> list[str]:
    """
    Break text into semantic chunks using NLTK
    """
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = []
    current_length = 0

    for sentence in sentences:
        # Add sentence length plus space
        sentence_length = len(sentence) + 1

        if current_length + sentence_length > chunk_size and current_chunk:
            chunks.append(' '.join(current_chunk))
            current_chunk = []
            current_length = 0

        current_chunk.append(sentence)
        current_length += sentence_length

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks

def process_document(file_content: bytes, filename: str) -> tuple[dict, str | list[str]]:
    """
    Process uploaded document based on file type
    Returns: (metadata, processed_content)
    """
    file_ext = Path(filename).suffix.lower()

    try:
        if file_ext in ['.csv', '.xlsx', '.xls']:
            if file_ext == '.csv':
                df = pd.read_csv(io.BytesIO(file_content))
            else:
                df = pd.read_excel(io.BytesIO(file_content))

            metadata = {
                "type": "spreadsheet",
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist()
            }
            return metadata, df.to_json(orient='records')

        elif file_ext == '.pdf':
            # Convert PDF to images
            images = convert_from_bytes(file_content)

            metadata = {
                "type": "pdf",
                "pages": len(images),
                "size": len(file_content)
            }

            # For now, just process the number of pages
            # In future, we can add OCR processing here
            text = f"PDF processed: {len(images)} pages"
            chunks = process_text_chunks(text)
            return metadata, chunks

        elif file_ext == '.docx':
            doc = docx.Document(io.BytesIO(file_content))
            full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            chunks = process_text_chunks(full_text)

            metadata = {
                "type": "docx",
                "paragraphs": len(doc.paragraphs),
                "chunks": len(chunks)
            }
            return metadata, chunks

        elif file_ext == '.txt':
            text = file_content.decode('utf-8')
            chunks = process_text_chunks(text)

            metadata = {
                "type": "text",
                "size": len(text),
                "chunks": len(chunks)
            }
            return metadata, chunks

        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    except Exception as e:
        raise Exception(f"Error processing document: {str(e)}")